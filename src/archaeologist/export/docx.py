"""DOCX export — convert markdown narrative to Word document."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def markdown_to_docx(markdown_text: str, output_path: Path, title: str = "Research Narrative"):
    """Convert a markdown narrative to a formatted DOCX file."""
    doc = Document()

    # Style setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    doc.add_heading(title, level=0)

    lines = markdown_text.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.left_indent = Inches(0.3)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            headers = [c.strip() for c in line.split('|') if c.strip()]
            i += 2  # skip header + separator

            rows = []
            while i < len(lines) and '|' in lines[i]:
                row = [c.strip() for c in lines[i].split('|') if c.strip()]
                rows.append(row)
                i += 1

            if headers:
                cols = len(headers)
                table = doc.add_table(rows=1 + len(rows), cols=cols, style='Light Grid Accent 1')
                for j, h in enumerate(headers):
                    table.rows[0].cells[j].text = h
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row[:cols]):
                        table.rows[r_idx + 1].cells[c_idx].text = cell
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^\d+\.\s+(.+)', line.strip())
        if num_match:
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, num_match.group(1))
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_formatted_text(p, line)
        i += 1

    doc.save(str(output_path))


def _add_formatted_text(paragraph, text: str):
    """Add text with basic inline formatting (bold, italic, code)."""
    # Process inline formatting
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
